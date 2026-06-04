import io
import re
from docx import Document

FIELD_KEYWORDS = {
    "product_name": ["наименование", "название продукта", "наименование продукта"],
    "product_category": ["категория", "вид продукта", "тип продукта"],
    "dosage_form": ["форма выпуска", "форма"],
    "net_weight": ["масса нетто", "нетто", "объём", "количество"],
    "serving_size": ["рекомендуемая доза", "разовая доза", "порция"],
    "servings_per_pack": ["количество порций", "порций в упаковке"],
    "composition": ["состав", "ингредиенты"],
    "bad_disclaimer": ["не является лекарственным", "бад", "биологически активная добавка"],
    "manufacturer_legal_name": ["производитель", "изготовитель"],
    "manufacturer_legal_address": ["юридический адрес"],
    "manufacturer_production_address": ["адрес производства", "место производства"],
    "manufacturer_complaints": ["контакты для претензий", "претензии", "контакты"],
    "sgr_number": ["свидетельство о государственной регистрации", "сгр", "номер сгр", "регистрационное удостоверение"],
    "sgr_date": ["дата сгр", "дата регистрации"],
    "tu_number": ["ту", "технические условия", "гост"],
    "shelf_life": ["срок годности", "срок хранения"],
    "storage_conditions": ["условия хранения", "хранить"],
}


def _find_field(paragraphs: list[str], keywords: list[str]) -> str | None:
    for i, para in enumerate(paragraphs):
        lower = para.lower()
        for kw in keywords:
            if kw in lower:
                # Value may be on the same line after colon, or on the next line
                if ":" in para:
                    value = para.split(":", 1)[1].strip()
                    if value:
                        return value
                if i + 1 < len(paragraphs):
                    return paragraphs[i + 1].strip()
    return None


def _parse_composition(paragraphs: list[str]) -> list[dict]:
    items = []
    in_composition = False
    for para in paragraphs:
        lower = para.lower()
        if any(kw in lower for kw in ["состав:", "ингредиенты:"]):
            in_composition = True
            text = para.split(":", 1)[1].strip() if ":" in para else ""
            if text:
                for part in re.split(r"[,;]", text):
                    part = part.strip()
                    if part:
                        match = re.match(r"(.+?)\s+([\d.,]+)\s*([а-яА-Яa-zA-Z/]+)?$", part)
                        if match:
                            items.append({"name": match.group(1).strip(), "amount": match.group(2), "unit": match.group(3) or ""})
                        else:
                            items.append({"name": part, "amount": "", "unit": ""})
            continue
        if in_composition:
            if para.strip() == "" or any(kw in lower for kw in list(FIELD_KEYWORDS.keys())):
                in_composition = False
                continue
            for part in re.split(r"[,;]", para):
                part = part.strip()
                if part:
                    match = re.match(r"(.+?)\s+([\d.,]+)\s*([а-яА-Яa-zA-Z/]+)?$", part)
                    if match:
                        items.append({"name": match.group(1).strip(), "amount": match.group(2), "unit": match.group(3) or ""})
                    else:
                        items.append({"name": part, "amount": "", "unit": ""})
    return items


def parse_pen_document(docx_bytes: bytes) -> dict:
    doc = Document(io.BytesIO(docx_bytes))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    paragraphs.append(text)

    result = {
        "product_name": _find_field(paragraphs, FIELD_KEYWORDS["product_name"]),
        "product_category": _find_field(paragraphs, FIELD_KEYWORDS["product_category"]),
        "dosage_form": _find_field(paragraphs, FIELD_KEYWORDS["dosage_form"]),
        "net_weight": _find_field(paragraphs, FIELD_KEYWORDS["net_weight"]),
        "serving_size": _find_field(paragraphs, FIELD_KEYWORDS["serving_size"]),
        "servings_per_pack": _find_field(paragraphs, FIELD_KEYWORDS["servings_per_pack"]),
        "composition": _parse_composition(paragraphs),
        "bad_disclaimer": _find_field(paragraphs, FIELD_KEYWORDS["bad_disclaimer"]),
        "manufacturer": {
            "legal_name": _find_field(paragraphs, FIELD_KEYWORDS["manufacturer_legal_name"]),
            "legal_address": _find_field(paragraphs, FIELD_KEYWORDS["manufacturer_legal_address"]),
            "production_address": _find_field(paragraphs, FIELD_KEYWORDS["manufacturer_production_address"]),
            "complaints_contact": _find_field(paragraphs, FIELD_KEYWORDS["manufacturer_complaints"]),
        },
        "sgr_number": _find_field(paragraphs, FIELD_KEYWORDS["sgr_number"]),
        "sgr_date": _find_field(paragraphs, FIELD_KEYWORDS["sgr_date"]),
        "tu_number": _find_field(paragraphs, FIELD_KEYWORDS["tu_number"]),
        "shelf_life": _find_field(paragraphs, FIELD_KEYWORDS["shelf_life"]),
        "storage_conditions": _find_field(paragraphs, FIELD_KEYWORDS["storage_conditions"]),
        "eaeu_mark_required": True,
    }
    return result
