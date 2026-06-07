import uuid
import io
from fastapi import APIRouter, Depends, File, Form, UploadFile, HTTPException
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from app.core.database import get_db
from app.core.deps import require_specialist
from app.models.users import User
from app.models.products import Product, ProductCategory
from app.models.files import Mockup, PenDocument
from app.services import file_service
from app.services.storage import storage_service
from app.schemas.files import (
    ProductCreate, ProductResponse, MockupResponse,
    PenDocumentResponse, ZipPreviewItem, ZipUploadRequest,
)

router = APIRouter(tags=["uploads"])


@router.post("/products", response_model=ProductResponse, status_code=201)
async def create_product(
    data: ProductCreate,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(require_specialist),
):
    product = Product(name=data.name, category=data.category, created_by=current_user.id)
    db.add(product)
    await db.commit()
    await db.refresh(product)
    return product


@router.get("/products", response_model=list[ProductResponse])
async def list_products(
    search: str = "",
    category: str = "",
    skip: int = 0,
    limit: int = 50,
    db: AsyncSession = Depends(get_db),
    _: User = Depends(require_specialist),
):
    q = select(Product)
    if search:
        q = q.where(Product.name.ilike(f"%{search}%"))
    if category:
        q = q.where(Product.category == category)
    result = await db.execute(q.offset(skip).limit(limit))
    return result.scalars().all()


@router.get("/products/{product_id}", response_model=ProductResponse)
async def get_product(
    product_id: uuid.UUID,
    db: AsyncSession = Depends(get_db),
    _: User = Depends(require_specialist),
):
    result = await db.execute(select(Product).where(Product.id == product_id))
    product = result.scalar_one_or_none()
    if not product:
        raise HTTPException(404, "Продукт не найден")
    return product


@router.get("/products/{product_id}/mockups", response_model=list[MockupResponse])
async def list_mockup_versions(
    product_id: uuid.UUID,
    db: AsyncSession = Depends(get_db),
    _: User = Depends(require_specialist),
):
    result = await db.execute(
        select(Mockup).where(Mockup.product_id == product_id).order_by(Mockup.version.desc())
    )
    return result.scalars().all()


@router.get("/products/{product_id}/pen", response_model=list[PenDocumentResponse])
async def list_pen_versions(
    product_id: uuid.UUID,
    db: AsyncSession = Depends(get_db),
    _: User = Depends(require_specialist),
):
    result = await db.execute(
        select(PenDocument).where(PenDocument.product_id == product_id).order_by(PenDocument.version.desc())
    )
    return result.scalars().all()


@router.post("/uploads/mockup", response_model=MockupResponse, status_code=201)
async def upload_mockup(
    file: UploadFile = File(...),
    product_id: uuid.UUID = Form(...),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(require_specialist),
):
    return await file_service.upload_mockup(db, product_id, file, current_user.id)


@router.post("/uploads/pen", response_model=PenDocumentResponse, status_code=201)
async def upload_pen(
    file: UploadFile = File(...),
    product_id: uuid.UUID = Form(...),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(require_specialist),
):
    return await file_service.upload_pen(db, product_id, file, current_user.id)


@router.post("/uploads/zip/preview", response_model=list[ZipPreviewItem])
async def preview_zip(
    file: UploadFile = File(...),
    _: User = Depends(require_specialist),
):
    if not file.filename or not file.filename.lower().endswith(".zip"):
        raise HTTPException(400, "Требуется ZIP-файл")
    content = await file.read()
    return await file_service.process_zip(content)


@router.post("/uploads/zip/confirm")
async def confirm_zip(
    file: UploadFile = File(...),
    category: str = Form("bad"),
    mapping_json: str = Form(...),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(require_specialist),
):
    import json
    content = await file.read()
    mapping = json.loads(mapping_json)
    cat = ProductCategory(category)
    return await file_service.confirm_zip_upload(db, content, mapping, current_user.id, cat)


@router.post("/uploads/reference")
async def upload_reference(
    file: UploadFile = File(...),
    _: User = Depends(require_specialist),
):
    """Extract plain text from a manual-review 'Замечание' file (PDF/DOCX/TXT)."""
    content = await file.read()
    name = (file.filename or "").lower()
    text = ""
    try:
        if name.endswith(".pdf"):
            import pdfplumber
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                text = "\n".join((p.extract_text() or "") for p in pdf.pages)
        elif name.endswith(".docx"):
            from docx import Document
            doc = Document(io.BytesIO(content))
            parts = [p.text for p in doc.paragraphs if p.text.strip()]
            for tbl in doc.tables:
                for row in tbl.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            parts.append(cell.text.strip())
            text = "\n".join(parts)
        else:  # txt / csv / fallback
            text = content.decode("utf-8", errors="ignore")
    except Exception as e:
        raise HTTPException(400, f"Не удалось прочитать файл замечаний: {e}")
    text = text.strip()
    if not text:
        raise HTTPException(400, "Файл замечаний пуст или текст не распознан")
    return {"text": text[:20000], "filename": file.filename}
