import io
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor


def generate_excel_report(task_id: str, issues: list[dict]) -> bytes:
    wb = Workbook()
    ws = wb.active
    ws.title = "Замечания"

    headers = ["№", "Тип", "Модуль", "Описание", "Предложение", "Статус"]
    header_fill = PatternFill(start_color="1F4E79", end_color="1F4E79", fill_type="solid")
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.font = Font(bold=True, color="FFFFFF", size=11)
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center")

    ws.column_dimensions["A"].width = 5
    ws.column_dimensions["B"].width = 12
    ws.column_dimensions["C"].width = 15
    ws.column_dimensions["D"].width = 50
    ws.column_dimensions["E"].width = 40
    ws.column_dimensions["F"].width = 12

    type_colors = {"error": "FFE0E0", "warning": "FFF3CD", "info": "E0F0FF"}
    for i, issue in enumerate(issues, 1):
        fill_color = type_colors.get(issue.get("type", "info"), "FFFFFF")
        row_fill = PatternFill(start_color=fill_color, end_color=fill_color, fill_type="solid")
        row = [
            i,
            issue.get("type", ""),
            issue.get("module", ""),
            issue.get("description", ""),
            issue.get("suggestion", ""),
            issue.get("status", "open"),
        ]
        for col, val in enumerate(row, 1):
            cell = ws.cell(row=i + 1, column=col, value=val)
            cell.fill = row_fill
            cell.alignment = Alignment(wrap_text=True)

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def generate_md_report(task_id: str, product_name: str, issues: list[dict]) -> bytes:
    errs = sum(1 for i in issues if i.get("type") == "error")
    warns = sum(1 for i in issues if i.get("type") == "warning")
    lines = [
        f"# Отчёт о проверке макета",
        f"- Продукт: {product_name}",
        f"- ID проверки: {task_id}",
        f"- Всего замечаний: {len(issues)} (ошибок: {errs}, предупреждений: {warns})",
        "",
        "| Тип | Модуль | Описание | Предложение |",
        "|---|---|---|---|",
    ]
    for i in issues:
        d = str(i.get("description", "")).replace("|", "\\|").replace("\n", " ")
        s = str(i.get("suggestion", "") or "").replace("|", "\\|").replace("\n", " ")
        lines.append(f"| {i.get('type','')} | {i.get('module','')} | {d} | {s} |")
    return ("\n".join(lines)).encode("utf-8")


def generate_batch_md(batch_id: str, items: list[dict]) -> bytes:
    """items: [{name, status, issues:[...]}] — consolidated batch report as Markdown."""
    lines = [f"# Пакетная проверка этикеток", f"Batch ID: {batch_id}", f"Этикеток: {len(items)}", "",
             "## Сводка", "", "| # | Этикетка | Статус | Замечаний |", "|---|---|---|---|"]
    for i, it in enumerate(items, 1):
        lines.append(f"| {i} | {it['name']} | {it['status']} | {len(it.get('issues') or [])} |")
    lines.append("")
    for it in items:
        lines.append(f"## {it['name']} — {it['status']} — замечаний: {len(it.get('issues') or [])}")
        lines.append("")
        lines.append("| Тип | Модуль | Описание | Предложение |")
        lines.append("|---|---|---|---|")
        for x in (it.get("issues") or []):
            d = str(x.get("description", "")).replace("|", "\\|").replace("\n", " ")
            s = str(x.get("suggestion", "") or "").replace("|", "\\|").replace("\n", " ")
            lines.append(f"| {x.get('type','')} | {x.get('module','')} | {d} | {s} |")
        lines.append("")
    return ("\n".join(lines)).encode("utf-8")


def generate_batch_word(batch_id: str, items: list[dict]) -> bytes:
    doc = DocxDocument()
    doc.add_heading("Пакетная проверка этикеток", 0)
    doc.add_paragraph(f"Batch ID: {batch_id}")
    doc.add_paragraph(f"Этикеток: {len(items)}")
    doc.add_heading("Сводка", 1)
    st = doc.add_table(rows=1, cols=4); st.style = "Table Grid"
    for i, h in enumerate(["#", "Этикетка", "Статус", "Замечаний"]):
        st.rows[0].cells[i].text = h; st.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
    for i, it in enumerate(items, 1):
        c = st.add_row().cells
        c[0].text = str(i); c[1].text = it["name"]; c[2].text = it["status"]; c[3].text = str(len(it.get("issues") or []))
    for it in items:
        doc.add_heading(f"{it['name']} — замечаний: {len(it.get('issues') or [])}", 1)
        t = doc.add_table(rows=1, cols=4); t.style = "Table Grid"
        for i, h in enumerate(["Тип", "Модуль", "Описание", "Предложение"]):
            t.rows[0].cells[i].text = h; t.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
        for x in (it.get("issues") or []):
            r = t.add_row().cells
            r[0].text = str(x.get("type", "")); r[1].text = str(x.get("module", ""))
            r[2].text = str(x.get("description", "")); r[3].text = str(x.get("suggestion", "") or "")
    buf = io.BytesIO(); doc.save(buf); return buf.getvalue()


def generate_word_report(task_id: str, product_name: str, issues: list[dict]) -> bytes:
    doc = DocxDocument()
    doc.add_heading(f"Отчёт о проверке макета", 0)
    doc.add_paragraph(f"Продукт: {product_name}")
    doc.add_paragraph(f"ID проверки: {task_id}")
    doc.add_paragraph(f"Всего замечаний: {len(issues)}")
    doc.add_paragraph(f"Ошибок: {sum(1 for i in issues if i.get('type')=='error')} | Предупреждений: {sum(1 for i in issues if i.get('type')=='warning')}")
    doc.add_heading("Список замечаний", 1)

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(["Тип", "Модуль", "Описание", "Предложение"]):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].font.bold = True

    for issue in issues:
        row = table.add_row().cells
        row[0].text = issue.get("type", "")
        row[1].text = issue.get("module", "")
        row[2].text = issue.get("description", "")
        row[3].text = issue.get("suggestion", "")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
